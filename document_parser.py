"""
document_parser.py
────────────────────────────────────────────────────────────────────────────
Extracts text / tabular data from uploaded documents and answers questions
about them using retrieval-augmented generation.

Supported formats: PDF, DOCX, TXT, CSV, XLSX.

Design notes
------------
* Large documents are NEVER passed to the LLM whole. Text is chunked and
  only the top-scoring chunks (via rag_utils.SimpleTfidfRetriever) are
  placed into the prompt.
* CSV / Excel files are summarised (shape, columns, quick stats) AND
  converted row-by-row into rich text so the same TF-IDF retrieval
  machinery can answer row-level questions ("who gave bad feedback?").
* Simple arithmetic questions (average / sum / count of a column) are
  additionally computed directly with pandas for accuracy, and the
  computed fact is injected into the context passed to the LLM.
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from typing import List, Optional

import pandas as pd

from rag_utils import chunk_text, SimpleTfidfRetriever

MAX_TEXT_CHARS_FOR_SUMMARY = 4000  # cap used only for the "quick preview"


@dataclass
class ParsedDocument:
    """Normalised result of parsing any supported document type."""
    kind: str                       # "text" or "table"
    filename: str
    raw_text: str = ""              # full extracted text (text-type files)
    dataframe: Optional[pd.DataFrame] = None   # table-type files
    summary: str = ""                # human-readable summary
    retriever: Optional[SimpleTfidfRetriever] = field(default=None, repr=False)
    row_texts: List[str] = field(default_factory=list)

    def build_retriever(self):
        if self.kind == "text":
            chunks = chunk_text(self.raw_text)
            self.retriever = SimpleTfidfRetriever(chunks)
        elif self.kind == "table":
            self.retriever = SimpleTfidfRetriever(self.row_texts)


# ─────────────────────────────────────────────────────────── EXTRACTORS

def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(pages).strip()


def _extract_docx(path: str) -> str:
    import docx
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()


def _row_to_text(row: dict, row_num: int) -> str:
    parts = [f"Row {row_num}:"]
    for col, val in row.items():
        sval = str(val).strip()
        if sval and sval.lower() not in ("nan", "none", ""):
            parts.append(f"{col}: {sval}")
    return "  ".join(parts)


def _dataframe_summary(df: pd.DataFrame, filename: str) -> str:
    lines = [
        f"File: {filename}",
        f"Rows: {len(df)}, Columns: {len(df.columns)}",
        f"Column names: {', '.join(str(c) for c in df.columns)}",
    ]
    numeric_cols = df.select_dtypes(include="number").columns.tolist()
    if numeric_cols:
        lines.append("Numeric column statistics:")
        stats = df[numeric_cols].describe().round(2)
        for col in numeric_cols:
            lines.append(
                f"  {col}: mean={stats.loc['mean', col]}, "
                f"min={stats.loc['min', col]}, max={stats.loc['max', col]}"
            )
    cat_cols = [c for c in df.columns if c not in numeric_cols]
    for col in cat_cols[:6]:
        try:
            top_vals = df[col].astype(str).value_counts().head(3)
            preview = ", ".join(f"{k} ({v})" for k, v in top_vals.items())
            lines.append(f"  {col} top values: {preview}")
        except Exception:
            continue
    return "\n".join(lines)


def _load_table(path: str, ext: str) -> pd.DataFrame:
    if ext == ".csv":
        return pd.read_csv(path)
    return pd.read_excel(path)


# ─────────────────────────────────────────────────────────── PUBLIC API

TEXT_EXTS = {".pdf", ".docx", ".txt"}
TABLE_EXTS = {".csv", ".xlsx", ".xls"}


def parse_document(path: str, filename: str) -> ParsedDocument:
    """
    Parse an uploaded document from disk and return a ParsedDocument ready
    for retrieval. Raises ValueError for unsupported / corrupt files.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        text = _extract_pdf(path)
        doc = ParsedDocument(kind="text", filename=filename, raw_text=text)
        doc.summary = (
            f"PDF document '{filename}' — extracted {len(text)} characters "
            f"of text."
        )

    elif ext == ".docx":
        text = _extract_docx(path)
        doc = ParsedDocument(kind="text", filename=filename, raw_text=text)
        doc.summary = (
            f"Word document '{filename}' — extracted {len(text)} characters "
            f"of text."
        )

    elif ext == ".txt":
        text = _extract_txt(path)
        doc = ParsedDocument(kind="text", filename=filename, raw_text=text)
        doc.summary = f"Text file '{filename}' — {len(text)} characters."

    elif ext in TABLE_EXTS:
        df = _load_table(path, ext)
        df.columns = [str(c).strip() for c in df.columns]
        row_texts = [
            _row_to_text(row.to_dict(), i + 1) for i, row in df.iterrows()
        ]
        summary = _dataframe_summary(df, filename)
        doc = ParsedDocument(
            kind="table", filename=filename, dataframe=df,
            summary=summary, row_texts=row_texts,
        )

    else:
        raise ValueError(f"Unsupported document type: {ext}")

    if not doc.raw_text and doc.kind == "text":
        raise ValueError(
            f"Could not extract any readable text from '{filename}'. "
            "The file may be empty, scanned, or corrupted."
        )

    doc.build_retriever()
    return doc


# ─────────────────────────────────────────────────────────── TABULAR Q&A

_AVG_PATTERN = re.compile(
    r"\b(average|mean|avg)\b.{0,40}?\b([a-zA-Z][a-zA-Z0-9 _]{1,30})\b", re.I
)
_COUNT_PATTERN = re.compile(r"\b(how many|count|number of|total)\b", re.I)


def compute_quick_stat(doc: ParsedDocument, query: str) -> Optional[str]:
    """
    Try to directly compute a simple statistic (average / count / filter)
    from the uploaded table using pandas, so numeric answers are accurate
    rather than left entirely to the LLM's judgement.

    Returns a fact string to inject into the LLM context, or None if no
    direct computation could be confidently made (falls back to pure RAG).
    """
    if doc.kind != "table" or doc.dataframe is None:
        return None
    df = doc.dataframe
    q_low = query.lower()

    # ── Average / mean of a numeric column ───────────────────────────────
    m = _AVG_PATTERN.search(q_low)
    if m:
        target = m.group(2).strip()
        for col in df.columns:
            if target in col.lower() or col.lower() in target:
                if pd.api.types.is_numeric_dtype(df[col]):
                    val = df[col].mean()
                    return f"Computed statistic: the average {col} is {round(val, 2)}."

    # ── Count / how many, optionally with a filter keyword ────────────────
    if _COUNT_PATTERN.search(q_low):
        filter_words = [w for w in re.findall(r"[a-zA-Z]+", q_low)
                         if len(w) > 3]
        best_col, best_mask, best_count = None, None, -1
        for col in df.columns:
            col_str = df[col].astype(str).str.lower()
            for w in filter_words:
                mask = col_str.str.contains(re.escape(w), na=False)
                cnt = int(mask.sum())
                if 0 < cnt < len(df) and cnt > best_count:
                    best_col, best_mask, best_count = col, mask, cnt
        if best_mask is not None:
            return (f"Computed statistic: {best_count} row(s) match, "
                    f"based on column '{best_col}'.")
        return f"Computed statistic: the file has {len(df)} total rows."

    return None


def retrieve_relevant_rows_or_chunks(doc: ParsedDocument, query: str,
                                       top_k: int = 5) -> List[str]:
    """Retrieve the most relevant chunks (text docs) or rows (tables)."""
    if doc.retriever is None:
        return []
    results = doc.retriever.retrieve(query, top_k=top_k)
    return [text for text, _score in results]
